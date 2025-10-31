"""
AI Contract Clause Builder - Assignment 2

import streamlit as st
import os
from io import BytesIO
from datetime import datetime
from docx import Document
from openai import OpenAI
import traceback


# ============================================================================
# 
# ============================================================================

def get_api_key_from_sidebar():
    """
    Get API key from sidebar
    
    Features:
    1. Priority: environment variable (for production)
    2. Fallback: user input (for local testing)
    
    Returns:
        str: OpenAI API key, None if not provided
    """
    st.sidebar.header("API Configuration")
    
    # Priority: environment variable (for deployment)
    env_key = os.getenv("OPENAI_API_KEY")
    if env_key:
        st.sidebar.success("Using API key from environment variable")
        return env_key
    
    # User manual input
    key = st.sidebar.text_input(
        "Enter your OpenAI API Key",
        type="password",
        key="api_key_input",
        help="Get API key: https://platform.openai.com/api-keys"
    )
    
    if key:
        st.sidebar.success("API key entered successfully")
    else:
        st.sidebar.warning("Please enter API key to continue")
    
    return key


def call_openai_chat(messages, api_key, model="gpt-4o-mini", temperature=0.2, max_tokens=1000):
    """
    OpenAI Chat API
    
    
    1. 
    2. 
    3. 
    4. AI Call times
    
    Args:
        messages: 
        api_key: API
        model: 
        temperature: 0-2
        max_tokens: token
        
    Returns:
        str: AI
    """
    try:
        # AI Call times
        if 'ai_call_count' not in st.session_state:
            st.session_state.ai_call_count = 0
        st.session_state.ai_call_count += 1
        
        # OpenAI API Configuration
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.openai.com/v1"
        )
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
            timeout=60  # 60
        )
        
        return resp.choices[0].message.content.strip()
        
    except Exception as e:
        error_type = type(e).__name__
        
        # 
        if "AuthenticationError" in error_type:
            st.error("**API Authentication Failed**\n\nPlease check your OpenAI API key.")
        elif "RateLimitError" in error_type:
            st.error(" **API**\n\n\n- \n- API\n- ")
        elif "timeout" in str(e).lower():
            st.error(" ****\n\nAPI")
        else:
            st.error(f" **API**\n\n: {error_type}\n: {str(e)}")
        
        # 
        with st.expander(" View Detailed Error InformationFor debugging"):
            st.code(traceback.format_exc())
        
        st.stop()  # 


def extract_text_from_uploaded_files(uploaded_files):
    """
    
    
    
    1. 
    2. 
    3. 
    
    Args:
        uploaded_files: Streamlit
        
    Returns:
        list: 
    """
    texts = []
    failed_files = []
    
    for uploaded in uploaded_files:
        try:
            if uploaded.name.lower().endswith(".docx"):
                # Word
                doc = Document(uploaded)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif uploaded.name.lower().endswith((".txt", ".md")):
                # 
                content = uploaded.getvalue()
                try:
                    text = content.decode("utf-8")
                except:
                    text = content.decode("latin-1")
            else:
                # 
                content = uploaded.getvalue()
                try:
                    text = content.decode("utf-8")
                except:
                    text = content.decode("latin-1")
            
            # 
            if text.strip():
                texts.append({"filename": uploaded.name, "text": text})
            else:
                failed_files.append(f"{uploaded.name} ()")
                
        except Exception as e:
            failed_files.append(f"{uploaded.name} (: {str(e)})")
    
    # 
    if failed_files:
        st.warning(f" \n" + "\n".join([f"- {f}" for f in failed_files]))
    
    return texts


def ai_enhanced_retrieve(texts, query, api_key, top_k=3):
    """
    AIRAG
    
    
    1. AI
    2. 
    3. Lab 2RAG
    
    Args:
        texts: 
        query: 
        api_key: API
        top_k: top k
        
    Returns:
        list: 
    """
    if not texts:
        return []
    
    # token
    doc_snippets = []
    for i, item in enumerate(texts):
        snippet = item['text'][:800]  # 800
        doc_snippets.append(f"[{i}] : {item['filename']}\n: {snippet}")
    
    combined_docs = "\n\n".join(doc_snippets)
    
    # AI
    retrieval_prompt = f"""

{query}


{combined_docs}

 {top_k} 

1. 
2. 

"No relevant documents"


X: 
"""
    
    try:
        retrieval_result = call_openai_chat(
            [
                {"role": "system", "content": ""},
                {"role": "user", "content": retrieval_prompt}
            ],
            api_key,
            temperature=0.1
        )
        
        # 
        relevant_indices = []
        for i in range(len(texts)):
            if f"{i}" in retrieval_result or f"[{i}]" in retrieval_result:
                relevant_indices.append(i)
        
        # 
        return [texts[i] for i in relevant_indices[:top_k]]
        
    except:
        # AI
        return simple_retrieve(texts, query, top_k)


def simple_retrieve(texts, query, top_k=3):
    """
    
    
    AI
    """
    q_words = set([w.lower() for w in query.split() if len(w) > 3])
    scored = []
    
    for item in texts:
        words = set([w.lower().strip('.,;:\n') for w in item['text'].split() if len(w) > 3])
        overlap = len(q_words & words)
        scored.append((overlap, item))
    
    scored.sort(reverse=True, key=lambda x: x[0])
    return [item for score, item in scored[:top_k] if score > 0]


def create_docx(clause_text, metadata):
    """
    Create a professional Word document with proper legal formatting
    
    Args:
        clause_text: The clause content
        metadata: Document metadata (timestamp, objective, etc.)
        
    Returns:
        BytesIO: Document binary stream
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title - Professional Legal Document Style
    title = doc.add_heading(level=1)
    title_run = title.add_run("CONTRACT CLAUSE")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Document Information Section
    info_heading = doc.add_heading("DOCUMENT INFORMATION", level=2)
    info_heading.runs[0].font.name = 'Times New Roman'
    info_heading.runs[0].font.size = Pt(12)
    
    # Metadata table-style layout
    info_para = doc.add_paragraph()
    info_para.add_run("Date Generated: ").bold = True
    info_para.add_run(f"{metadata.get('timestamp', 'N/A')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Drafting Objective: ").bold = True
    info_para.add_run(f"{metadata.get('objective', 'N/A')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Jurisdiction: ").bold = True
    info_para.add_run(f"{metadata.get('jurisdiction', 'Not specified')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Drafting Style: ").bold = True
    info_para.add_run(f"{metadata.get('style', 'N/A')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Clause Content Section
    clause_heading = doc.add_heading("CLAUSE PROVISIONS", level=2)
    clause_heading.runs[0].font.name = 'Times New Roman'
    clause_heading.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing before clause
    
    # Add clause text with proper formatting (remove all markdown and LaTeX artifacts)
    import re
    
    # Function to clean LaTeX formatting
    def clean_latex(text):
        # Remove display math delimiters \[ \]
        text = re.sub(r'\\\[(.*?)\\\]', r'\1', text, flags=re.DOTALL)
        # Remove inline math delimiters \( \)
        text = re.sub(r'\\\((.*?)\\\)', r'\1', text, flags=re.DOTALL)
        # Remove $$ $$ delimiters
        text = re.sub(r'\$\$(.*?)\$\$', r'\1', text, flags=re.DOTALL)
        # Remove $ $ delimiters
        text = re.sub(r'\$(.*?)\$', r'\1', text)
        # Replace LaTeX commands with readable text
        text = text.replace(r'\text{', '').replace('}', '')
        text = text.replace(r'\times', '×')
        text = text.replace(r'\%', '%')
        # Replace \frac{a}{b} with (a / b)
        text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1 / \2)', text)
        # Remove remaining backslashes
        text = text.replace('\\', '')
        return text
    
    # Split into lines and process each
    lines = clause_text.split('\n')
    
    for line in lines:
        if not line.strip():
            continue  # Skip empty lines
            
        # Remove markdown formatting
        clean_line = line.replace("**", "").replace("*", "")
        # Remove markdown headers but keep the text
        clean_line = re.sub(r'^#{1,6}\s+', '', clean_line)
        # Clean LaTeX formatting
        clean_line = clean_latex(clean_line)
        
        # Check if this looks like a header (was markdown header or is short and looks like title)
        if line.startswith('#'):
            # This was a markdown header - make it a sub-heading in Word
            heading_para = doc.add_paragraph(clean_line)
            heading_para.runs[0].bold = True
            heading_para.runs[0].font.name = 'Times New Roman'
            heading_para.runs[0].font.size = Pt(11)
            heading_para.paragraph_format.space_before = Pt(6)
        else:
            # Regular paragraph
            clause_para = doc.add_paragraph(clean_line)
            clause_para.paragraph_format.line_spacing = 1.15
            for run in clause_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing after clause
    
    # Professional Disclaimer
    disclaimer_heading = doc.add_heading("DISCLAIMER", level=2)
    disclaimer_heading.runs[0].font.name = 'Times New Roman'
    disclaimer_heading.runs[0].font.size = Pt(12)
    
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "This document has been generated using artificial intelligence technology "
        "and is provided for informational purposes only. It does not constitute "
        "legal advice, and should not be relied upon as such. Users should consult "
        "with qualified legal professionals before using any content from this document "
        "in actual legal agreements or contracts. The creators and distributors of this "
        "tool disclaim all liability for any damages arising from the use of this document."
    )
    disclaimer.paragraph_format.line_spacing = 1.15
    for run in disclaimer.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.italic = True
    
    # Save to memory
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio


# ============================================================================
# Streamlit 
# ============================================================================

# Page configuration
st.set_page_config(
    page_title="AI Contract Clause Builder",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Title and description
st.title("AI Contract Clause Builder")
st.markdown("""
An AI-powered legal tool for automated contract clause drafting and review.

### Key Features:
- Multi-step AI analysis: 7 discrete steps ensuring clause quality
- Document support: Upload reference documents (optional)
- Automated review: Multiple rounds of AI refinement
- Word export: Generate professional Word documents
- Legal expertise: Based on legal professional knowledge

### Disclaimer:
This tool generates content for informational purposes only and does not constitute legal advice. 
Consult qualified legal professionals before formal use.
""")

# Usage examples
with st.expander("Click to view usage examples and recommendations"):
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        **Sample Objectives:**
        - "Limit liability for indirect, special, or consequential damages"
        - "Specify liquidated damages calculation method and cap"
        - "Define intellectual property ownership and licensing scope"
        - "Establish confidentiality obligations and duration"
        - "Determine dispute resolution method and jurisdiction"
        """)
    
    with col2:
        st.markdown("""
        **Usage Recommendations:**
        - Provide clear and specific objective descriptions
        - Upload reference documents if available (improves quality)
        - Choose appropriate drafting style
        - Recommend at least 2 automated reviews
        - Carefully review generated content
        """)

st.markdown("---")

# ============================================================================
#  - 
# ============================================================================

# API
api_key = get_api_key_from_sidebar()

st.sidebar.markdown("---")
st.sidebar.header("Input Parameters")

# 1. Clause objective
objective = st.sidebar.text_area(
    "Clause Drafting Objective",
    height=120,
    placeholder="e.g., Limit liability for indirect damages to 20% of contract amount",
    key="objective_input",
    help="Clearly describe the purpose and requirements for the clause"
)

# 2. Jurisdiction
jurisdiction = st.sidebar.text_input(
    "Jurisdiction (Optional)",
    placeholder="e.g., Mainland China / Hong Kong / Singapore",
    key="jurisdiction_input",
    help="Different jurisdictions may have different legal requirements"
)

# 3. File upload
uploaded_files = st.sidebar.file_uploader(
    "Upload Reference Documents (Optional)",
    accept_multiple_files=True,
    type=['txt', 'docx', 'md'],
    key="upload_input",
    help="Upload relevant contracts, cases, or reference documents (TXT, DOCX, MD)"
)

if uploaded_files:
    st.sidebar.success(f"Uploaded {len(uploaded_files)} file(s)")

# 4. Drafting style
firm_style = st.sidebar.selectbox(
    "Drafting Style",
    ["Plain English", "Legal Formal", "Balanced (Legal but Readable)"],
    index=2,
    key="style_input",
    help="Select the language style for the clause"
)

# 5. Review iterations
num_refinements = st.sidebar.slider(
    "Number of Automated Reviews",
    min_value=1,
    max_value=4,
    value=2,
    key="refinement_slider",
    help="Number of automated review and refinement iterations (higher = better quality, longer time)"
)

st.sidebar.markdown("---")

# Run button
run_button = st.sidebar.button(
    "Generate Clause",
    type="primary",
    use_container_width=True,
    key="run_button"
)

# Display statistics
if 'ai_call_count' in st.session_state:
    st.sidebar.info(f"AI Calls in Current Session: {st.session_state.ai_call_count}")

# ============================================================================
# 
# ============================================================================

if run_button:
    # AI Call
    st.session_state.ai_call_count = 0
    
    # 
    if not api_key:
        st.error(" Please enter your OpenAI API key in the sidebar")
        st.stop()
    
    if not objective or len(objective.strip()) < 10:
        st.error(" Please provide a clear clause objectiveat least10 characters")
        st.stop()
    
    # 
    progress_container = st.container()
    
    with progress_container:
        # 
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Step7Step + num_refinementsStep
        total_steps = 7 + num_refinements
        current_step = 0
        
        # ====================================================================
        # Step 1: 
        # ====================================================================
        current_step += 1
        status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Analyzing your objective...")
        progress_bar.progress(current_step / total_steps)
        
        st.markdown("## Step 1: Objective Analysis")
        st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Interpret drafting objective")
        
        with st.spinner("..."):
            interpretation = call_openai_chat(
                [
                    {"role": "system", "content": ""},
                    {"role": "user", "content": f"""5

****: {objective}

****: {jurisdiction or 'Not specified'}

****: {firm_style}


1. 
2. 
3. 
4. 
5. 
"""}
                ],
                api_key
            )
        
        st.success(" Objective Analysis")
        st.markdown(interpretation)
        
        # ====================================================================
        # Step 2: 
        # ====================================================================
        current_step += 1
        status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Processing documents...")
        progress_bar.progress(current_step / total_steps)
        
        st.markdown("## Step 2: Document Analysis and Legal Research")
        
        texts = extract_text_from_uploaded_files(uploaded_files) if uploaded_files else []
        
        if texts:
            # Uploaded Documents
            st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Summarize uploaded documents")
            
            combined_preview = "\n\n".join([
                f"--- {t['filename']} ---\n{t['text'][:1500]}"  # 1500
                for t in texts
            ])
            
            with st.expander(" View Uploaded Document Preview"):
                st.text_area("Document Content Preview", combined_preview, height=200, key="docs_preview")
            
            with st.spinner("Summarizing documents..."):
                docs_summary = call_openai_chat(
                    [
                        {"role": "system", "content": ""},
                        {"role": "user", "content": f"""

****: {objective}

****:
{combined_preview}


1. 
2. 
3. 
4. 
"""}
                    ],
                    api_key
                )
            
            st.success(" Document summary completed")
            st.markdown(docs_summary)
            
            # AIRAG
            st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Intelligent retrieval of relevant segments")
            
            with st.spinner("Retrieving relevant content..."):
                retrieved = ai_enhanced_retrieve(texts, objective, api_key, top_k=3)
            
            if retrieved:
                st.success(f" Found {len(retrieved)} relevant document segment(s)")
                with st.expander(" View Retrieved Relevant Segments"):
                    for r in retrieved:
                        st.markdown(f"** {r['filename']}**")
                        st.code(r['text'][:500], language="text")
            else:
                st.info("ℹ Found")
        else:
            # Uploaded DocumentsConduct legal background researchAI Call times
            st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Conduct legal background research")
            
            with st.spinner("Researching relevant legal background..."):
                legal_research = call_openai_chat(
                    [
                        {"role": "system", "content": ""},
                        {"role": "user", "content": f"""

****: {objective}
****: {jurisdiction or ''}


1. 
2. 
3. 
4. 
"""}
                    ],
                    api_key
                )
            
            st.success(" Legal research completed")
            st.markdown(legal_research)
            
            docs_summary = f"(Uploaded Documents)\n\n\n{legal_research}"
            retrieved = []
        
        # ====================================================================
        # Step 3: 
        # ====================================================================
        current_step += 1
        status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Analyzing constraints...")
        progress_bar.progress(current_step / total_steps)
        
        st.markdown("## Step 3: Constraints and Risk Analysis")
        st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Analyze constraints and legal risks")
        
        # 
        evidence_block = "\n\n".join([
            f"From {r['filename']}\n{r['text'][:1000]}"
            for r in retrieved
        ]) if retrieved else "(No relevant documents)"
        
        with st.spinner("Analyzing constraints..."):
            constraints = call_openai_chat(
                [
                    {"role": "system", "content": ""},
                    {"role": "user", "content": f"""

****: {objective}

****: {jurisdiction or ''}

****: {docs_summary}

****: {evidence_block}



**A. **
3-5

**B. **
2-3

**C. **
3-5

**D. **

"""}
                ],
                api_key
            )
        
        st.success(" Analysis completed")
        st.markdown(constraints)
        
        # ====================================================================
        # Step 4: Draft Initial Clause
        # ====================================================================
        current_step += 1
        status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Drafting clause...")
        progress_bar.progress(current_step / total_steps)
        
        st.markdown("## Step 4: Draft Initial Clause")
        st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Draft initial clause version")
        
        with st.spinner("Drafting clause..."):
            initial_clause = call_openai_chat(
                [
                    {"role": "system", "content": "You are an experienced contract lawyer. CRITICAL: Use PLAIN TEXT only - NO LaTeX (no \\frac, \\text, \\[, \\]). For math use: (A / B) format"},
                    {"role": "user", "content": f"""

****: {objective}

****: {constraints}

****: {firm_style}

****:
1. 
2. 
3. 
4. 






Drafting Notes
• 1: ...
• 2: ...
• 3: ...
"""}
                ],
                api_key,
                max_tokens=1500
            )
        
        st.success(" Initial clause drafting completed")
        
        # Parse and display initial clause
        if "Drafting Notes" in initial_clause:
            parts = initial_clause.split("Drafting Notes")
            clause_part = parts[0].strip()
            explanation_part = parts[1].strip() if len(parts) > 1 else ""
            
            st.markdown("### Initial Clause Version")
            st.code(clause_part, language="text")
            
            if explanation_part:
                with st.expander("View Drafting Notes"):
                    st.markdown(explanation_part)
        else:
            clause_part = initial_clause
            st.markdown("### Initial Clause Version")
            st.code(initial_clause, language="text")
        
        # ====================================================================
        # Step 5: Review and Refinement
        # ====================================================================
        st.markdown("## Step 5: Review and Refinement")
        
        current_clause = clause_part if "Drafting Notes" in initial_clause else initial_clause
        
        for i in range(num_refinements):
            current_step += 1
            status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Conducting review {i+1}...")
            progress_bar.progress(current_step / total_steps)
            
            st.markdown(f"### Review Round {i+1}")
            st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Review and refine clause")
            
            with st.spinner(f"Conducting review  {i+1} ..."):
                review = call_openai_chat(
                    [
                        {"role": "system", "content": "You are a professional contract lawyer conducting a thorough review of legal clauses."},
                        {"role": "user", "content": f"""
Please review and refine the following contract clause.

**Drafting Objective**: {objective}

**Current Clause**:
{current_clause}

**Review Requirements**:
1. Check legal completeness and accuracy
2. Improve language clarity and precision
3. Ensure enforceability under relevant jurisdiction
4. Add necessary qualifications or exceptions
5. Optimize structure and readability

**Output Format** (IMPORTANT - Follow this exact format):

[Revised Clause]
(Complete revised clause text here)

[Revision Notes]
- First improvement description (what was changed and why)
- Second improvement description  
- Third improvement description
(Use bullet points with dashes, NOT numbered lists like "1.", "2." etc.)
"""}
                    ],
                    api_key,
                    max_tokens=1500
                )
            
            # Parse review result - handle both [Revised Clause] and Revised Clause formats
            if "Revised Clause" in review or "[Revised Clause]" in review:
                # Try to split by Revision Notes (with or without brackets)
                if "[Revision Notes]" in review:
                    parts = review.split("[Revision Notes]")
                elif "Revision Notes" in review:
                    parts = review.split("Revision Notes")
                else:
                    parts = [review, ""]
                
                # Extract revised clause (remove format markers)
                revised_clause = parts[0].replace("[Revised Clause]", "").replace("Revised Clause", "").strip()
                # Remove any remaining brackets or parenthetical instructions
                revised_clause = revised_clause.replace("(Complete revised clause text here)", "").strip()
                
                changes = parts[1].strip() if len(parts) > 1 else ""
                # Remove instruction text from changes
                changes = changes.replace("(Use bullet points with dashes, NOT numbered lists like \"1.\", \"2.\" etc.)", "").strip()
                
                st.success(f"Review {i+1} completed")
                
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.markdown("**Revised Clause**")
                    st.code(revised_clause, language="text")
                
                with col2:
                    st.markdown("**Revision Notes**")
                    if changes:
                        st.markdown(changes)
                    else:
                        st.info("No specific changes documented")
                
                current_clause = revised_clause
            else:
                # Fallback: use entire review result as the revised clause
                st.success(f"Review {i+1} completed")
                st.markdown("**Revised Clause**")
                st.code(review, language="text")
                current_clause = review
        
        # ====================================================================
        # Step 6: Final Version
        # ====================================================================
        current_step += 1
        status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Generating final version...")
        progress_bar.progress(current_step / total_steps)
        
        st.markdown("## Step 6: Final Version")
        st.success(" review completed")
        
        st.markdown("###  Final Clause")
        st.code(current_clause, language="text")
        
        # Word
        metadata = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "objective": objective,
            "jurisdiction": jurisdiction or "Not specified",
            "style": firm_style,
            "ai_calls": st.session_state.ai_call_count
        }
        
        docx_bio = create_docx(current_clause, metadata)
        
        st.download_button(
            label=" DownloadWord",
            data=docx_bio,
            file_name=f"AI_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
        # ====================================================================
        # Step 7: Quality Assessment
        # ====================================================================
        current_step += 1
        status_text.info(f"⏳ Progress: {current_step}/{total_steps} - Assessing quality...")
        progress_bar.progress(current_step / total_steps)
        
        st.markdown("## Step 7: Quality Assessment")
        st.info(f" AI Call #{st.session_state.ai_call_count + 1}: Assess clause quality")
        
        with st.spinner("Assess clause quality..."):
            evaluation = call_openai_chat(
                [
                    {"role": "system", "content": "You are a senior legal expert conducting quality assessment of contract clauses."},
                    {"role": "user", "content": f"""
Please assess the quality of the following contract clause.

**Drafting Objective**: {objective}

**Final Clause**:
{current_clause}

**Assessment Requirements**:
Please score the clause from the following 10 dimensions (10 points each, total 100 points):

1. Objective Achievement - Does it fulfill the drafting objective?
2. Legal Validity - Is it legally sound and compliant?
3. Language Clarity - Is the wording clear and unambiguous?
4. Logical Rigor - Is the structure logical and coherent?
5. Enforceability - Can it be effectively enforced?
6. Risk Control - Does it adequately address potential risks?
7. Professionalism - Does it meet professional legal standards?
8. Completeness - Are all necessary elements included?
9. Applicability - Is it practical and applicable?
10. Overall Quality - Overall assessment

**Output Format** (IMPORTANT - Follow exactly):

[Scoring]
1. Objective Achievement: X/10
2. Legal Validity: X/10
3. Language Clarity: X/10
4. Logical Rigor: X/10
5. Enforceability: X/10
6. Risk Control: X/10
7. Professionalism: X/10
8. Completeness: X/10
9. Applicability: X/10
10. Overall Quality: X/10
Total Score: XX/100

[Strengths]
• Strength 1
• Strength 2
• Strength 3

[Areas for Improvement]
• Suggestion 1
• Suggestion 2
• Suggestion 3
"""}
                ],
                api_key
            )
        
        st.success(" Quality Assessment")
        st.markdown(evaluation)
        
        # ====================================================================
        # 
        # ====================================================================
        status_text.success(f" StepTotal AI calls: {st.session_state.ai_call_count}  times")
        progress_bar.progress(1.0)
        
        # 
        st.markdown("---")
        st.markdown("###  Generation Statistics")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("AI Call times", st.session_state.ai_call_count)
        
        with col2:
            st.metric("Step", 7)
        
        with col3:
            st.metric(" times", num_refinements)
        
        with col4:
            st.metric("Uploaded Documents", len(texts) if texts else 0)
        
        # 
        st.balloons()

# ============================================================================
# 
# ============================================================================

st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; font-size: 0.9em;'>
    <p> AI | </p>
    <p> </p>
    <p> OpenAI API</p>
</div>
""", unsafe_allow_html=True)
